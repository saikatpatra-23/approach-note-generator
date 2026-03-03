# -*- coding: utf-8 -*-
"""
doc_generator.py
Builds the Approach Note Word document -- functional format matching the
Tata Motors / Tata Technologies standard template style.
"""
from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from prompts import SECTIONS

# Colour palette
TATA_BLUE  = RGBColor(0x00, 0x47, 0xAB)
HDR_FILL   = "0047AB"   # table header fill
ALT_FILL   = "DCE6F1"   # alternating row / label fill
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY  = RGBColor(0x44, 0x44, 0x44)

# Standard app list for Impact Analysis
IMPACT_APPS = [
    ("OLTP (Siebel CRM)",           ""),
    ("OLAP (Siebel Analytics)",      ""),
    ("Mobility / iRA App",           ""),
    ("BIP Reports",                  ""),
    ("EIM (Batch Jobs)",             ""),
    ("SAP / IS Auto",                ""),
    ("DMS (Dealer Mgmt System)",     ""),
    ("CTI (Telephony Integration)",  ""),
    ("Portal / External App",        ""),
]

WORK_TYPES = [
    "SRF Changes",
    "Interface / API Changes",
    "EIM Changes",
    "BIP Report Changes",
    "Smartscript Changes",
    "Mobility Changes",
    "Data Retrofitment",
    "Other",
]


# == Low-level XML helpers ====================================================

def _fill_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_white(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = WHITE
        if not para.runs and para.text:
            run = para.add_run(para.text)
            para.clear()
            run.bold = True
            run.font.color.rgb = WHITE


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


# == Public entry point ========================================================

def build_approach_note(
    template_bytes: bytes,
    cover_details: dict,
    sections_dict: dict,
) -> bytes:
    tmpl = Document(io.BytesIO(template_bytes))
    doc  = Document()

    # Copy page margins from template
    ts = tmpl.sections[0]
    ds = doc.sections[0]
    ds.page_width       = ts.page_width
    ds.page_height      = ts.page_height
    ds.left_margin      = ts.left_margin
    ds.right_margin     = ts.right_margin
    ds.top_margin       = ts.top_margin
    ds.bottom_margin    = ts.bottom_margin

    # Title
    _add_title(doc, cover_details)

    # Cover table
    doc.add_paragraph()
    _add_cover_table(doc, cover_details)

    # Document history
    doc.add_paragraph()
    _add_heading(doc, "Document History", level=2)
    _add_history_table(doc, cover_details)

    # Table of contents placeholder
    doc.add_paragraph()
    _add_heading(doc, "Table of Contents", level=2)
    toc_para = doc.add_paragraph()
    toc_run  = toc_para.add_run("[Auto-generated from section headings]")
    toc_run.italic = True
    toc_run.font.color.rgb = DARK_GREY

    # Content sections
    for key, display_name in SECTIONS.items():
        content = sections_dict.get(key, "")
        doc.add_paragraph()
        _add_heading(doc, display_name, level=1)

        if key == "impact_analysis":
            _add_impact_analysis(doc, content)
        elif key == "open_items":
            _add_open_items_table(doc, content)
        else:
            _add_text_section(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# == Section writers ===========================================================

def _add_title(doc: Document, cover: dict) -> None:
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"Approach Note - {cover.get('cr_number', '')}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = TATA_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(cover.get("summary", ""))
    sub_run.font.size = Pt(11)
    sub_run.italic = True
    sub_run.font.color.rgb = DARK_GREY


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    run  = para.add_run(text)
    run.font.color.rgb = TATA_BLUE


def _add_cover_table(doc: Document, cover: dict) -> None:
    _add_heading(doc, "Document Details", level=2)

    fields = [
        ("CRQ Number",      cover.get("cr_number", "")),
        ("Source",          "BRD"),
        ("Change Type",     cover.get("change_type", "")),
        ("Timeline",        cover.get("timeline", "")),
        ("Project",         cover.get("project", "")),
        ("Application",     cover.get("application", "")),
        ("Module",          cover.get("module", "")),
        ("Business Unit",   cover.get("business_unit", "")),
        ("Complexity",      cover.get("complexity", "")),
        ("BRM",             cover.get("brm_name", "")),
        ("BPO",             cover.get("bpo_name", "")),
        ("BA",              cover.get("ba_name", "")),
        ("Readership",      "Functional Team, Development Team, Testing Team, Business Stakeholders"),
        ("Summary",         cover.get("summary", "")),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    col_widths = [Inches(2.0), Inches(4.5)]

    for i, (label, value) in enumerate(fields):
        row = table.rows[i]
        for j, width in enumerate(col_widths):
            row.cells[j].width = width
        row.cells[0].text = label
        row.cells[1].text = str(value)
        _fill_cell(row.cells[0], "DCE6F1")
        _bold_cell(row.cells[0])


def _add_history_table(doc: Document, cover: dict) -> None:
    headers = ["Ver.", "Date", "Author(s)", "Comments / Changes"]
    table   = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
        _fill_cell(hdr.cells[i], HDR_FILL)
        _bold_white(hdr.cells[i])

    data = table.rows[1]
    data.cells[0].text = "1.0"
    data.cells[1].text = date.today().strftime("%d-%b-%Y")
    data.cells[2].text = cover.get("ba_name", "")
    data.cells[3].text = "Draft"


def _add_text_section(doc: Document, content: Any) -> None:
    """Render a plain text / bullet section."""
    if not content:
        doc.add_paragraph("(To be updated)")
        return

    text = content if isinstance(content, str) else str(content)
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(line[2:])
        else:
            doc.add_paragraph(line)


def _add_impact_analysis(doc: Document, content: Any) -> None:
    """
    Render Impact Analysis with two sub-tables:
    A. Application Impact
    B. Technical Work Types
    """
    # ── A. Application Impact ──────────────────────────────────────────────
    _add_heading(doc, "A. Impacted Applications", level=2)

    # Build lookup from Claude's output
    app_lookup: dict[str, tuple[str, str]] = {}
    if isinstance(content, dict) and "applications" in content:
        for item in content["applications"]:
            key = str(item.get("app", "")).lower()
            app_lookup[key] = (
                str(item.get("impacted", "N")),
                str(item.get("remarks", "")),
            )

    app_headers = ["Sr. No.", "Application", "Impacted (Y/N)", "Remarks"]
    app_table   = doc.add_table(rows=1 + len(IMPACT_APPS), cols=4)
    app_table.style = "Table Grid"

    hdr = app_table.rows[0]
    for i, h in enumerate(app_headers):
        hdr.cells[i].text = h
        _fill_cell(hdr.cells[i], HDR_FILL)
        _bold_white(hdr.cells[i])

    for r_idx, (app_name, _) in enumerate(IMPACT_APPS, 1):
        row = app_table.rows[r_idx]
        # fuzzy match
        match_key = None
        for k in app_lookup:
            if k and any(word in app_name.lower() for word in k.split()):
                match_key = k
                break

        impacted, remarks = app_lookup.get(match_key or "", ("N", ""))
        row.cells[0].text = str(r_idx)
        row.cells[1].text = app_name
        row.cells[2].text = impacted
        row.cells[3].text = remarks

        if impacted.upper() == "Y":
            _fill_cell(row.cells[2], "C6EFCE")
        elif "PARTIAL" in impacted.upper():
            _fill_cell(row.cells[2], "FFEB9C")
        else:
            _fill_cell(row.cells[2], "FFCCCC")

    # ── B. Technical Work Types ────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "B. Technical Work Type", level=2)

    wt_lookup: dict[str, str] = {}
    if isinstance(content, dict) and "work_types" in content:
        for item in content["work_types"]:
            key = str(item.get("type", "")).lower()
            wt_lookup[key] = str(item.get("applicable", "N"))

    wt_headers = ["Sr. No.", "Technical Work Type", "Applicable (Y/N)"]
    wt_table   = doc.add_table(rows=1 + len(WORK_TYPES), cols=3)
    wt_table.style = "Table Grid"

    hdr2 = wt_table.rows[0]
    for i, h in enumerate(wt_headers):
        hdr2.cells[i].text = h
        _fill_cell(hdr2.cells[i], HDR_FILL)
        _bold_white(hdr2.cells[i])

    for r_idx, wt_name in enumerate(WORK_TYPES, 1):
        row = wt_table.rows[r_idx]
        match_key = None
        for k in wt_lookup:
            if k and any(word in wt_name.lower() for word in k.split()):
                match_key = k
                break
        applicable = wt_lookup.get(match_key or "", "N")
        row.cells[0].text = str(r_idx)
        row.cells[1].text = wt_name
        row.cells[2].text = applicable
        _fill_cell(row.cells[2], "C6EFCE" if applicable.upper() == "Y" else "F2F2F2")


def _add_open_items_table(doc: Document, content: Any) -> None:
    if not content or (isinstance(content, list) and len(content) == 0):
        doc.add_paragraph("No open items at this stage.")
        return

    if isinstance(content, list) and content and isinstance(content[0], dict):
        headers = ["S. No.", "Open Item", "Owner", "Status"]
        table   = doc.add_table(rows=1 + len(content), cols=4)
        table.style = "Table Grid"

        hdr = table.rows[0]
        for i, h in enumerate(headers):
            hdr.cells[i].text = h
            _fill_cell(hdr.cells[i], HDR_FILL)
            _bold_white(hdr.cells[i])

        for r_idx, item in enumerate(content, 1):
            row = table.rows[r_idx]
            row.cells[0].text = str(item.get("sno", r_idx))
            row.cells[1].text = str(item.get("item", ""))
            row.cells[2].text = str(item.get("owner", ""))
            row.cells[3].text = str(item.get("status", "Open"))
    else:
        _add_text_section(doc, content)
