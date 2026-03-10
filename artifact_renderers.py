from __future__ import annotations

import io
import json
from datetime import date
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from schemas import GeneratedArtifact, WorkspaceSnapshot

BRAND_BLUE = RGBColor(0x11, 0x4B, 0x8C)
HEADER_FILL = "114B8C"
LABEL_FILL = "D9E6F2"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def render_artifact(
    artifact: GeneratedArtifact,
    workspace: WorkspaceSnapshot,
    template_bytes: bytes | None = None,
) -> tuple[bytes, str, str]:
    if artifact.renderer == "excel":
        payload = build_effort_workbook(artifact, workspace)
        return payload, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx"
    if artifact.renderer == "svg":
        payload = build_diagram_svg(artifact, workspace)
        return payload, "image/svg+xml", ".svg"
    payload = build_word_artifact(artifact, workspace, template_bytes)
    return payload, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"


def build_word_artifact(
    artifact: GeneratedArtifact,
    workspace: WorkspaceSnapshot,
    template_bytes: bytes | None = None,
) -> bytes:
    template_doc = Document(io.BytesIO(template_bytes)) if template_bytes else None
    doc = Document()
    if template_doc:
        src = template_doc.sections[0]
        dst = doc.sections[0]
        dst.page_width = src.page_width
        dst.page_height = src.page_height
        dst.left_margin = src.left_margin
        dst.right_margin = src.right_margin
        dst.top_margin = src.top_margin
        dst.bottom_margin = src.bottom_margin

    _add_title(doc, artifact.title)
    _add_workspace_details(doc, workspace, artifact)
    doc.add_paragraph()

    for key, value in artifact.payload.items():
        if key == "title":
            continue
        _add_heading(doc, key.replace("_", " ").title(), level=1)
        if key == "impact_analysis" and isinstance(value, dict):
            _add_impact_analysis(doc, value)
        elif key == "open_items" and isinstance(value, list):
            _add_open_items(doc, value)
        elif isinstance(value, list):
            _add_list_or_table(doc, value)
        elif isinstance(value, dict):
            _add_dict_section(doc, value)
        else:
            _add_text(doc, str(value))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_effort_workbook(artifact: GeneratedArtifact, workspace: WorkspaceSnapshot) -> bytes:
    wb = Workbook()
    summary_sheet = wb.active
    summary_sheet.title = "Summary"
    summary_sheet["A1"] = artifact.title
    summary_sheet["A1"].font = Font(size=14, bold=True)
    summary_sheet["A3"] = "Workspace"
    summary_sheet["B3"] = workspace.name
    summary_sheet["A4"] = "Application"
    summary_sheet["B4"] = workspace.application_name
    summary_sheet["A5"] = "Module"
    summary_sheet["B5"] = workspace.module_name
    summary_sheet["A6"] = "Generated On"
    summary_sheet["B6"] = date.today().isoformat()
    summary_sheet["A8"] = "Estimate Summary"
    summary_sheet["A9"] = artifact.payload.get("estimate_summary", "")

    assumptions = wb.create_sheet("Assumptions")
    assumptions["A1"] = "Assumptions"
    assumptions["A1"].font = Font(size=12, bold=True)
    for idx, item in enumerate(artifact.payload.get("assumptions", []), start=2):
        assumptions[f"A{idx}"] = item

    estimate_sheet = wb.create_sheet("Estimate")
    headers = ["Workstream", "Activity", "Role", "Effort Days", "Remarks"]
    for col, header in enumerate(headers, start=1):
        cell = estimate_sheet.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(fill_type="solid", fgColor=HEADER_FILL)

    for row_idx, line in enumerate(artifact.payload.get("work_breakdown", []), start=2):
        estimate_sheet.cell(row=row_idx, column=1, value=line.get("workstream", ""))
        estimate_sheet.cell(row=row_idx, column=2, value=line.get("activity", ""))
        estimate_sheet.cell(row=row_idx, column=3, value=line.get("role", ""))
        estimate_sheet.cell(row=row_idx, column=4, value=line.get("effort_days", 0))
        estimate_sheet.cell(row=row_idx, column=5, value=line.get("remarks", ""))

    totals = wb.create_sheet("Totals")
    totals["A1"] = "Totals"
    totals["A1"].font = Font(size=12, bold=True)
    totals["A3"] = "Total Effort Days"
    totals["B3"] = artifact.payload.get("totals", {}).get("total_effort_days", 0)
    totals["A4"] = "Recommended Team Shape"
    totals["B4"] = artifact.payload.get("totals", {}).get("recommended_team_shape", "")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def build_diagram_svg(artifact: GeneratedArtifact, workspace: WorkspaceSnapshot) -> bytes:
    nodes = artifact.payload.get("nodes", [])
    edges = artifact.payload.get("edges", [])
    width = 1160
    height = max(420, 140 + len(nodes) * 110)
    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#F8FBFD"/>',
        f'<text x="40" y="42" font-size="28" font-family="Segoe UI, Arial" font-weight="700" fill="#114B8C">{escape(artifact.payload.get("title", artifact.title))}</text>',
        f'<text x="40" y="72" font-size="14" font-family="Segoe UI, Arial" fill="#4A5C6A">{escape(artifact.payload.get("objective", workspace.summary)[:140])}</text>',
        '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#46607A" /></marker></defs>',
    ]

    positions: dict[str, tuple[int, int]] = {}
    columns = 3 if len(nodes) > 4 else 2
    box_width = 250
    box_height = 72
    x_gap = 80
    y_gap = 95
    base_x = 60
    base_y = 120

    for index, node in enumerate(nodes):
        col = index % columns
        row = index // columns
        x = base_x + col * (box_width + x_gap)
        y = base_y + row * (box_height + y_gap)
        positions[node["id"]] = (x, y)
        svg_parts.append(
            f'<rect x="{x}" y="{y}" rx="16" ry="16" width="{box_width}" height="{box_height}" fill="#FFFFFF" stroke="#114B8C" stroke-width="2" />'
        )
        svg_parts.append(
            f'<text x="{x + 16}" y="{y + 32}" font-size="17" font-family="Segoe UI, Arial" font-weight="600" fill="#123A5B">{escape(node.get("label", ""))}</text>'
        )
        if node.get("category"):
            svg_parts.append(
                f'<text x="{x + 16}" y="{y + 55}" font-size="12" font-family="Segoe UI, Arial" fill="#567086">{escape(node.get("category", "").title())}</text>'
            )

    for edge in edges:
        source = positions.get(edge.get("source"))
        target = positions.get(edge.get("target"))
        if not source or not target:
            continue
        start_x = source[0] + box_width
        start_y = source[1] + box_height / 2
        end_x = target[0]
        end_y = target[1] + box_height / 2
        svg_parts.append(
            f'<line x1="{start_x}" y1="{start_y}" x2="{end_x}" y2="{end_y}" stroke="#46607A" stroke-width="2.4" marker-end="url(#arrow)" />'
        )
        if edge.get("label"):
            label_x = (start_x + end_x) / 2
            label_y = (start_y + end_y) / 2 - 8
            svg_parts.append(
                f'<text x="{label_x}" y="{label_y}" text-anchor="middle" font-size="12" font-family="Segoe UI, Arial" fill="#46607A">{escape(edge.get("label", ""))}</text>'
            )

    notes = artifact.payload.get("notes", [])
    if notes:
        notes_y = height - 80
        svg_parts.append(
            f'<text x="40" y="{notes_y}" font-size="15" font-family="Segoe UI, Arial" font-weight="600" fill="#114B8C">Notes</text>'
        )
        for idx, note in enumerate(notes[:4], start=1):
            svg_parts.append(
                f'<text x="40" y="{notes_y + idx * 22}" font-size="13" font-family="Segoe UI, Arial" fill="#4A5C6A">- {escape(note)}</text>'
            )

    svg_parts.append("</svg>")
    return "".join(svg_parts).encode("utf-8")


def _add_title(doc: Document, title: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_BLUE


def _add_workspace_details(doc: Document, workspace: WorkspaceSnapshot, artifact: GeneratedArtifact) -> None:
    _add_heading(doc, "Document Details", level=2)
    fields = [
        ("Workspace", workspace.name),
        ("Artifact", artifact.artifact_type.replace("_", " ").title()),
        ("Application", workspace.application_name),
        ("Module", workspace.module_name),
        ("Audience", workspace.audience),
        ("Domain", workspace.inferred_domain),
        ("Generated On", date.today().strftime("%d-%b-%Y")),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"
    for idx, (label, value) in enumerate(fields):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value or "-"
        _fill_cell(table.rows[idx].cells[0], LABEL_FILL)
        _bold_cell(table.rows[idx].cells[0])


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph(style=f"Heading {level}")
    para.add_run(text).font.color.rgb = BRAND_BLUE


def _add_text(doc: Document, content: str) -> None:
    if not content:
        doc.add_paragraph("(To be updated)")
        return
    for line in content.split("\n"):
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith("- "):
            doc.add_paragraph(clean[2:], style="List Bullet")
        else:
            doc.add_paragraph(clean)


def _add_impact_analysis(doc: Document, content: dict) -> None:
    _add_heading(doc, "Application Areas", level=2)
    app_table = doc.add_table(rows=1, cols=3)
    app_table.style = "Table Grid"
    headers = ["Area", "Impact", "Remarks"]
    for idx, header in enumerate(headers):
        app_table.rows[0].cells[idx].text = header
        _fill_cell(app_table.rows[0].cells[idx], HEADER_FILL)
        _bold_white(app_table.rows[0].cells[idx])

    for item in content.get("application_areas", []):
        row = app_table.add_row()
        row.cells[0].text = item.get("area", "")
        row.cells[1].text = item.get("impact", "")
        row.cells[2].text = item.get("remarks", "")

    doc.add_paragraph()
    _add_heading(doc, "Change Dimensions", level=2)
    dim_table = doc.add_table(rows=1, cols=2)
    dim_table.style = "Table Grid"
    dim_table.rows[0].cells[0].text = "Dimension"
    dim_table.rows[0].cells[1].text = "Applicable"
    for cell in dim_table.rows[0].cells:
        _fill_cell(cell, HEADER_FILL)
        _bold_white(cell)

    for item in content.get("change_dimensions", []):
        row = dim_table.add_row()
        row.cells[0].text = item.get("dimension", "")
        row.cells[1].text = item.get("applicable", "")


def _add_open_items(doc: Document, items: list[dict]) -> None:
    if not items:
        doc.add_paragraph("No open items at this stage.")
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["S. No.", "Open Item", "Owner", "Status"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        _fill_cell(table.rows[0].cells[idx], HEADER_FILL)
        _bold_white(table.rows[0].cells[idx])
    for item in items:
        row = table.add_row()
        row.cells[0].text = str(item.get("sno", ""))
        row.cells[1].text = item.get("item", "")
        row.cells[2].text = item.get("owner", "")
        row.cells[3].text = item.get("status", "")


def _add_list_or_table(doc: Document, values: list) -> None:
    if not values:
        doc.add_paragraph("(None)")
        return
    if isinstance(values[0], dict):
        headers = list(values[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header.replace("_", " ").title()
            _fill_cell(table.rows[0].cells[idx], HEADER_FILL)
            _bold_white(table.rows[0].cells[idx])
        for item in values:
            row = table.add_row()
            for idx, header in enumerate(headers):
                row.cells[idx].text = json.dumps(item.get(header), ensure_ascii=True) if isinstance(item.get(header), (dict, list)) else str(item.get(header, ""))
        return
    for item in values:
        doc.add_paragraph(str(item), style="List Bullet")


def _add_dict_section(doc: Document, values: dict) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in values.items():
        row = table.add_row()
        row.cells[0].text = key.replace("_", " ").title()
        row.cells[1].text = json.dumps(value, indent=2) if isinstance(value, (dict, list)) else str(value)
        _fill_cell(row.cells[0], LABEL_FILL)
        _bold_cell(row.cells[0])


def _fill_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _bold_white(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = WHITE
        if not para.runs and para.text:
            run = para.add_run(para.text)
            para.clear()
            run.bold = True
            run.font.color.rgb = WHITE


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
